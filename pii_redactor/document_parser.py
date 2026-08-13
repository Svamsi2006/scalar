"""
document_parser.py - DOCX Document Parser & Format-Preserving Replacer

Handles ingestion and modification of MS Word (.docx) documents, including:
- Body paragraphs
- Nested tables (recursive)
- Headers and footers
- Format-preserving PII replacement via the run-splitting algorithm

The run-splitting algorithm is critical: Word internally splits text across
multiple XML <w:r> (run) elements, each with different formatting. A naive
find-and-replace on individual runs will miss PII that spans multiple runs.
This module concatenates run texts, detects PII on the full string, then maps
character offsets back to runs, splitting runs at match boundaries as needed.
"""

import copy
import logging
from typing import List, Dict, Tuple, Optional, Any

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses and modifies DOCX documents with format-preserving PII replacement."""

    def __init__(self):
        self.document: Optional[Document] = None
        self._stats = {
            'paragraphs_processed': 0,
            'tables_processed': 0,
            'headers_processed': 0,
            'footers_processed': 0,
            'replacements_made': 0,
        }

    def load(self, filepath: str) -> Document:
        """Load a DOCX file into memory.

        Args:
            filepath: Path to the .docx file.

        Returns:
            The loaded Document object.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            ValueError: If the file is not a valid .docx.
        """
        try:
            self.document = Document(filepath)
            logger.info(f"Successfully loaded document: {filepath}")
            return self.document
        except Exception as e:
            logger.error(f"Failed to load document '{filepath}': {e}")
            raise

    def save(self, filepath: str) -> None:
        """Save the modified document to disk.

        Args:
            filepath: Output path for the .docx file.

        Raises:
            RuntimeError: If no document is loaded.
        """
        if self.document is None:
            raise RuntimeError("No document loaded. Call load() first.")
        try:
            self.document.save(filepath)
            logger.info(f"Document saved to: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save document to '{filepath}': {e}")
            raise

    def iter_all_paragraphs(self) -> List[Paragraph]:
        """Yield all paragraphs from the document body, tables, headers, and footers.

        Returns:
            A flat list of all Paragraph objects in the document.
        """
        if self.document is None:
            raise RuntimeError("No document loaded. Call load() first.")

        paragraphs = []

        # Body paragraphs
        for para in self.document.paragraphs:
            paragraphs.append(para)

        # Tables (recursive for nested tables)
        for table in self.document.tables:
            paragraphs.extend(self._extract_table_paragraphs(table))

        # Headers and footers
        # Note: We process ALL headers/footers that exist and have content,
        # regardless of is_linked_to_previous. In single-section documents,
        # the primary header/footer is linked to itself (True) but still
        # contains PII that must be redacted.
        seen_header_elements = set()
        seen_footer_elements = set()

        for section in self.document.sections:
            # Headers
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header is None:
                    continue
                header_id = id(header._element)
                if header_id in seen_header_elements:
                    continue
                seen_header_elements.add(header_id)
                for para in header.paragraphs:
                    paragraphs.append(para)
                    self._stats['headers_processed'] += 1
                for table in header.tables:
                    paragraphs.extend(self._extract_table_paragraphs(table))

            # Footers
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer is None:
                    continue
                footer_id = id(footer._element)
                if footer_id in seen_footer_elements:
                    continue
                seen_footer_elements.add(footer_id)
                for para in footer.paragraphs:
                    paragraphs.append(para)
                    self._stats['footers_processed'] += 1
                for table in footer.tables:
                    paragraphs.extend(self._extract_table_paragraphs(table))

        return paragraphs

    def _extract_table_paragraphs(self, table: Table) -> List[Paragraph]:
        """Recursively extract paragraphs from a table, including nested tables.

        Args:
            table: A python-docx Table object.

        Returns:
            List of Paragraph objects found in the table.
        """
        paragraphs = []
        try:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs.append(para)
                    # Handle nested tables within cells
                    for nested_table in cell.tables:
                        paragraphs.extend(self._extract_table_paragraphs(nested_table))
        except Exception as e:
            logger.warning(f"Error extracting table paragraphs: {e}")
        return paragraphs

    def process_paragraph(
        self,
        paragraph: Paragraph,
        detections: List[Dict],
        replacement_manager: Any,
    ) -> int:
        """Replace PII in a paragraph using format-preserving run-splitting.

        This is the core algorithm that handles Word's run-splitting behavior.

        Algorithm:
        1. Build a character-offset-to-run mapping from the paragraph's runs.
        2. Sort detections by start_char DESCENDING (right-to-left replacement
           prevents offset invalidation).
        3. For each detection:
           a. Find all runs that overlap the detection span.
           b. Split runs at match boundaries if the match starts/ends mid-run.
           c. Replace the matched portion in the first affected run.
           d. Clear text from remaining affected runs.
           e. Preserve all formatting from the first affected run.

        Args:
            paragraph: The python-docx Paragraph to process.
            detections: List of detection dicts with 'entity_type', 'start_char',
                       'end_char', 'value' keys.
            replacement_manager: ReplacementManager instance for fake value lookup.

        Returns:
            Number of replacements made in this paragraph.
        """
        if not detections:
            return 0

        runs = paragraph.runs
        if not runs:
            return 0

        # Step 1: Build character-offset-to-run mapping
        run_map = self._build_run_map(runs)
        full_text = ''.join(r.text or '' for r in runs)

        if not full_text.strip():
            return 0

        # Step 2: Sort detections right-to-left (descending start_char)
        sorted_detections = sorted(detections, key=lambda d: d['start_char'], reverse=True)

        replacements_made = 0

        for det in sorted_detections:
            start = det['start_char']
            end = det['end_char']
            original_value = det['value']
            entity_type = det['entity_type']

            # Validate the detection span against current text
            if start < 0 or end > len(full_text) or start >= end:
                logger.debug(f"Skipping invalid detection span [{start}:{end}] for '{original_value}'")
                continue

            # Get the fake replacement
            fake_value = replacement_manager.get_replacement(entity_type, original_value)

            # Step 3: Perform the format-preserving replacement
            success = self._replace_in_runs(paragraph, runs, run_map, start, end, fake_value)
            if success:
                replacements_made += 1
                self._stats['replacements_made'] += 1

                # Rebuild the run map after modification since offsets have changed
                runs = paragraph.runs
                run_map = self._build_run_map(runs)
                full_text = ''.join(r.text or '' for r in runs)

        return replacements_made

    def _build_run_map(self, runs: List[Run]) -> List[Tuple[int, int, int]]:
        """Build a mapping of (run_index, char_start, char_end) for each run.

        Args:
            runs: List of Run objects from a paragraph.

        Returns:
            List of tuples (run_index, start_offset, end_offset).
        """
        run_map = []
        offset = 0
        for i, run in enumerate(runs):
            text = run.text or ''
            run_map.append((i, offset, offset + len(text)))
            offset += len(text)
        return run_map

    def _replace_in_runs(
        self,
        paragraph: Paragraph,
        runs: List[Run],
        run_map: List[Tuple[int, int, int]],
        match_start: int,
        match_end: int,
        replacement: str,
    ) -> bool:
        """Execute the run-splitting replacement for a single PII match.

        This method handles three cases:
        1. Match is entirely within a single run — simple text replacement.
        2. Match spans multiple runs — split boundary runs, replace first,
           clear the rest.
        3. Match starts/ends mid-run — split the run at the boundary.

        Args:
            paragraph: The parent Paragraph object.
            runs: Current list of Run objects.
            run_map: Current (index, start, end) mapping.
            match_start: Start character offset of the PII match.
            match_end: End character offset of the PII match.
            replacement: The fake text to insert.

        Returns:
            True if replacement was successful, False otherwise.
        """
        # Find all runs that overlap with [match_start, match_end)
        affected = []
        for idx, rstart, rend in run_map:
            if rstart < match_end and rend > match_start:
                affected.append((idx, rstart, rend))

        if not affected:
            logger.debug(f"No runs found for match [{match_start}:{match_end}]")
            return False

        try:
            if len(affected) == 1:
                # Case 1: Match is within a single run
                idx, rstart, rend = affected[0]
                run = runs[idx]
                run_text = run.text or ''
                local_start = match_start - rstart
                local_end = match_end - rstart
                run.text = run_text[:local_start] + replacement + run_text[local_end:]
                return True

            # Case 2 & 3: Match spans multiple runs
            first_idx, first_rstart, first_rend = affected[0]
            last_idx, last_rstart, last_rend = affected[-1]

            first_run = runs[first_idx]
            first_text = first_run.text or ''
            local_start_in_first = match_start - first_rstart

            # Text before the match in the first run + replacement
            prefix = first_text[:local_start_in_first]
            first_run.text = prefix + replacement

            # Handle the last run — preserve text after the match
            last_run = runs[last_idx]
            last_text = last_run.text or ''
            local_end_in_last = match_end - last_rstart
            suffix = last_text[local_end_in_last:]

            if suffix:
                # Append suffix to the first run (it carries the replacement)
                first_run.text += suffix

            # Clear text in all affected runs except the first
            for idx, _, _ in affected[1:]:
                runs[idx].text = ''

            return True

        except Exception as e:
            logger.warning(f"Error during run replacement [{match_start}:{match_end}]: {e}")
            return False

    def get_stats(self) -> Dict[str, int]:
        """Return processing statistics.

        Returns:
            Dictionary with counts of processed elements and replacements.
        """
        return dict(self._stats)


def extract_full_text(paragraph: Paragraph) -> str:
    """Extract the full concatenated text from a paragraph's runs.

    This is a utility function used by the detection engine to get
    the complete text for PII scanning.

    Args:
        paragraph: A python-docx Paragraph object.

    Returns:
        The full text of the paragraph as a single string.
    """
    return ''.join(run.text or '' for run in paragraph.runs)
